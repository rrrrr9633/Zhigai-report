import base64
import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "zhigai-report-server" / "scripts" / "generate_report.py"
TEMPLATE_PATH = (
    ROOT
    / "zhigai-report-server"
    / "assets"
    / "智改数转调研表-20260507v1.0.docx"
)
SPEC = importlib.util.spec_from_file_location("generate_report", MODULE_PATH)
generate_report = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(generate_report)

# 1×1 PNG。测试只关心 Word 段落锚点，不依赖额外图像库。
PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


class TemplateImageSlotTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.image_path = Path(self.temp_dir.name) / "slot.png"
        self.image_path.write_bytes(PNG_BYTES)

    def generator(self):
        return generate_report.ZhigaiReportGenerator(str(TEMPLATE_PATH))

    def test_equipment_descriptions_use_fangsong_gb2312_body_format(self) -> None:
        generator = self.generator()
        descriptions = {
            "整体情况描述": "设备整体情况",
            "设备数控化、自动化水平概述": "数控化说明",
            "设备自动化控制方面": "自动控制说明",
            "设备互联方面": "设备互联说明",
            "设备监控与决策方面": "监控决策说明",
            "设备清单": [],
        }

        generator.fill_shebei_yingyong(descriptions)

        paragraphs = generator.doc.paragraphs
        overall = next(paragraph for paragraph in paragraphs if paragraph.text == "设备整体情况")
        bodies = [
            next(paragraph for paragraph in paragraphs if paragraph.text == text)
            for text in ("数控化说明", "自动控制说明", "设备互联说明", "监控决策说明")
        ]
        self.assertAlmostEqual(overall.paragraph_format.first_line_indent.pt, 32.0)
        for paragraph in [overall, *bodies]:
            self.assertAlmostEqual(paragraph.paragraph_format.line_spacing.pt, 30.0)
            self.assertEqual(len(paragraph.runs), 1)
            run = paragraph.runs[0]
            self.assertEqual(run.font.name, "仿宋_GB2312")
            self.assertEqual(run._r.rPr.rFonts.get(qn("w:eastAsia")), "仿宋_GB2312")
            self.assertAlmostEqual(run.font.size.pt, 16.0)
            self.assertFalse(run.bold)
        for paragraph in bodies:
            self.assertIsNotNone(paragraph._p.pPr.numPr)
            self.assertEqual(paragraph._p.pPr.numPr.numId.val, 3)
            self.assertEqual(paragraph.paragraph_format.space_before.pt, 0.0)
            self.assertEqual(paragraph.paragraph_format.space_after.pt, 0.0)

    def test_replaced_equipment_and_software_tables_fill_correct_columns(self) -> None:
        generator = self.generator()
        generator.fill_shebei_yingyong(
            {
                "整体情况描述": "设备整体情况",
                "设备清单": [
                    {
                        "关键装备种类": "智能加工装备",
                        "名称": "设备A",
                        "规格/型号": "型号A",
                        "单台设备价格": "12",
                        "系统状态": "在用",
                        "供应商": "供应商A",
                        "原厂商所在地": "省内",
                    }
                ],
            }
        )
        generator.fill_ruanjian_yingyong(
            "软件整体情况",
            [
                {
                    "关键软件种类": "经营管理类",
                    "名称": "ERP",
                    "规格/型号": "V1",
                    "部署方式": "本地部署",
                    "单套软件价格": "20",
                    "系统状态": "在用",
                    "供应商": "供应商B",
                    "原厂商所在地": "国内/省外",
                }
            ],
        )

        equipment = generator._table_by_first_header("关键装备种类")
        software = generator._table_by_first_header("关键软件种类")
        self.assertEqual(len(equipment.columns), 7)
        self.assertEqual(len(software.columns), 8)
        self.assertEqual(
            [cell.text for cell in equipment.rows[1].cells],
            ["智能加工装备", "设备A", "型号A", "12", "在用", "供应商A", "省内"],
        )
        self.assertEqual(
            [cell.text for cell in software.rows[1].cells],
            ["经营管理类", "ERP", "V1", "本地部署", "20", "在用", "供应商B", "国内/省外"],
        )
        for table in (equipment, software):
            header_properties = table.rows[0]._tr.get_or_add_trPr()
            self.assertIsNotNone(header_properties.find(qn("w:tblHeader")))
            for row in table.rows:
                self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))

    def test_other_section_uses_unnumbered_template_anchor(self) -> None:
        generator = self.generator()

        generator.fill_qita_qingkuang("无其他特殊情况。")

        paragraphs = generator.doc.paragraphs
        heading_index = next(
            index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "其他情况"
        )
        self.assertEqual(paragraphs[heading_index + 1].text.strip(), "无其他特殊情况。")
        self.assertEqual(paragraphs[heading_index].style.name, "Heading 3")

    def test_maturity_image_replaces_exact_blank_slot(self) -> None:
        generator = self.generator()

        generator._insert_maturity_score_image(str(self.image_path))

        paragraphs = generator.doc.paragraphs
        heading_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "智能制造成熟度现状评测具体情况"
        )
        self.assertTrue(has_drawing(paragraphs[heading_index + 1]))
        self.assertEqual(
            paragraphs[heading_index + 2].text.strip(),
            "智能制造成熟度现状分析",
        )

    def test_enterprise_images_follow_new_overview_field_slots(self) -> None:
        generator = self.generator()

        generator.fill_qiye_zongti(
            {
                "总体概述": "总体概述正文",
                "主营业务": "主营业务正文",
                "主导产品": ["产品A", "产品B"],
                "主要工艺与设备": [
                    {"名称": "产品A工艺", "流程": "工序A→工序B；使用设备A"}
                ],
                "竞争优势": "竞争优势正文",
            },
            {
                "第一部分": [str(self.image_path)],
                "第二部分": [str(self.image_path)],
                "第三部分": [str(self.image_path)],
            },
        )

        paragraphs = generator.doc.paragraphs
        overview_heading = next(
            index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "（二）企业总体情况"
        )
        next_heading = next(
            index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "（三）企业智改数转情况"
        )
        section = paragraphs[overview_heading + 1 : next_heading]
        texts = [paragraph.text.strip() for paragraph in section]
        image_indexes = [index for index, paragraph in enumerate(section) if has_drawing(paragraph)]

        self.assertEqual(len(image_indexes), 3)
        self.assertEqual(texts[image_indexes[0] - 1], "总体概述正文")
        self.assertEqual(texts[image_indexes[1] - 1], "产品B")
        self.assertEqual(texts[image_indexes[2] - 1], "工序A→工序B；使用设备A")
        self.assertLess(image_indexes[0], texts.index("主营业务：主营业务正文"))
        self.assertLess(image_indexes[1], texts.index("主要工艺与设备：产品A工艺"))
        self.assertLess(image_indexes[2], texts.index("竞争优势：竞争优势正文"))
        self.assertFalse(any(text.startswith("内容要求：包括") for text in texts))

    def test_architecture_image_replaces_template_placeholder_after_layer_content(self) -> None:
        generator = self.generator()
        architecture_names = [
            "基础层（基础设施保障）",
            "设备与控制层（设备数采集成）",
            "平台层（平台支撑）",
            "管理应用层（企业管理应用）",
            "智能决策层（BI+AI 决策）",
        ]

        generator.fill_zhigai_jianshe_fangan(
            {
                "企业名称": "测试企业",
                "建设思路": "建设思路",
                "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                "总体方案架构": [
                    {"层级": name, "内容": f"{name}内容"} for name in architecture_names
                ],
                "总体技术架构图": str(self.image_path),
                "建设内容描述": [],
                "项目预期成效": {},
                "具体改造项目": [],
            }
        )

        paragraphs = generator.doc.paragraphs
        final_layer_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "智能决策层（BI+AI 决策）"
        )
        caption_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "总体技术架构图"
        )
        construction_content_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "建设内容描述"
        )
        slot_image_indexes = [
            index
            for index in range(final_layer_index + 1, caption_index)
            if has_drawing(paragraphs[index])
        ]
        self.assertEqual(len(slot_image_indexes), 1)
        self.assertLess(final_layer_index, slot_image_indexes[0])
        self.assertLess(slot_image_indexes[0], caption_index)
        self.assertLess(caption_index, construction_content_index)
        final_layer_content_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "智能决策层（BI+AI 决策）内容"
        )
        self.assertLess(final_layer_index, final_layer_content_index)
        self.assertLess(final_layer_content_index, slot_image_indexes[0])
        self.assertAlmostEqual(generator.doc.inline_shapes[-1].width / 914400, 6.0, places=2)

    def test_architecture_without_replacement_retains_template_image_and_caption(self) -> None:
        generator = self.generator()
        generator.fill_zhigai_jianshe_fangan(
            {
                "企业名称": "测试企业",
                "建设思路": "建设思路",
                "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                "总体方案架构": [
                    {"层级": name}
                    for name in (
                        "基础层（基础设施保障）",
                        "设备与控制层（设备数采集成）",
                        "平台层（平台支撑）",
                        "管理应用层（企业管理应用）",
                        "智能决策层（BI+AI 决策）",
                    )
                ],
                "建设内容描述": [],
                "项目预期成效": {},
                "具体改造项目": [],
            }
        )

        paragraphs = generator.doc.paragraphs
        caption_index = next(
            index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "总体技术架构图"
        )
        self.assertTrue(has_drawing(paragraphs[caption_index - 1]))
        self.assertEqual(len(generator.doc.inline_shapes), 1)

    def test_digital_overview_replaces_instruction_before_equipment_section(self) -> None:
        generator = self.generator()

        generator.fill_digital_overview("企业数字化智能化现状概述正文。")

        paragraphs = generator.doc.paragraphs
        overview_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text == "企业数字化智能化现状概述正文。"
        )
        equipment_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "数字化智能化设备应用情况"
        )
        self.assertLess(overview_index, equipment_index)
        self.assertFalse(
            any(paragraph.text.startswith("内容要求：XXXXXXXXX") for paragraph in paragraphs)
        )

    def test_pain_analysis_overview_and_details_use_first_line_indent(self) -> None:
        generator = self.generator()
        generator.fill_tongdian_fenxi(
            {
                "总体概述": "痛点总体概述",
                "痛点列表": [
                    {"名称": "生产痛点", "内容": "生产痛点正文"},
                    {"名称": "质量痛点", "内容": "质量痛点正文"},
                ],
            }
        )

        paragraphs = generator.doc.paragraphs
        for text in ("痛点总体概述", "生产痛点正文", "质量痛点正文"):
            paragraph = next(item for item in paragraphs if item.text == text)
            self.assertAlmostEqual(paragraph.paragraph_format.first_line_indent.pt, 32.0)
            self.assertAlmostEqual(paragraph.paragraph_format.line_spacing.pt, 30.0)

    def test_project_effects_indent_plain_body_but_not_bullets(self) -> None:
        generator = self.generator()
        generator.fill_zhigai_jianshe_fangan(
            {
                "企业名称": "测试企业",
                "建设思路": "建设思路",
                "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                "总体方案架构": [],
                "建设内容描述": [],
                "项目预期成效": {
                    "项目概括": "覆盖设备、系统和数据建设",
                    "智能制造成熟度提升": {
                        "智能制造成熟度评分": "成熟度提升说明",
                        "智能工厂梯度培育等级": "梯度等级说明",
                        "核心能力": "核心能力说明",
                    },
                    "效率提升方面": "效率提升正文",
                    "成本降低方面": "成本降低正文",
                    "合规与质量成效": "合规质量正文",
                    "运营与战略成效": "运营战略正文",
                },
                "具体改造项目": [],
            }
        )

        paragraphs = generator.doc.paragraphs
        intro = next(item for item in paragraphs if item.text.startswith("本项目覆盖设备"))
        self.assertAlmostEqual(intro.paragraph_format.first_line_indent.pt, 32.0)
        for text in ("效率提升正文", "成本降低正文", "合规质量正文", "运营战略正文"):
            paragraph = next(item for item in paragraphs if item.text == text)
            self.assertAlmostEqual(paragraph.paragraph_format.first_line_indent.pt, 32.0)
        bullet = next(item for item in paragraphs if item.text == "智能制造成熟度评分：成熟度提升说明")
        self.assertIsNotNone(bullet._p.pPr.numPr)
        bullet_indent = bullet.paragraph_format.first_line_indent
        self.assertTrue(bullet_indent is None or abs(bullet_indent.pt - 32.0) > 0.01)

    def test_three_year_plan_fields_take_priority_and_cover_36_months(self) -> None:
        generator = self.generator()
        generator.fill_zhigai_jianshe_fangan(
            {
                "企业名称": "测试企业",
                "建设思路": "建设思路",
                "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                "三年规划": [
                    {"阶段": "第一阶段", "时间": "0-12个月", "核心任务": "任务一", "关键成果": "成果一"},
                    {"阶段": "第二阶段", "时间": "13-24个月", "核心任务": "任务二", "关键成果": "成果二"},
                    {"阶段": "第三阶段", "时间": "25-36个月", "核心任务": "任务三", "关键成果": "成果三"},
                ],
                "两年规划": [
                    {"阶段": "旧阶段", "时间": "0-24个月", "核心任务": "旧任务", "关键成果": "旧成果"}
                ],
                "总体方案架构": [],
                "建设内容描述": [],
                "三年实施计划": [
                    {"阶段": "第一阶段", "周期": "0-12个月", "核心任务": "实施一"},
                    {"阶段": "第二阶段", "周期": "13-24个月", "核心任务": "实施二"},
                    {"阶段": "第三阶段", "周期": "25-36个月", "核心任务": "实施三"},
                ],
                "两年实施计划": [
                    {"阶段": "旧阶段", "周期": "0-24个月", "核心任务": "旧实施"}
                ],
                "项目预期成效": {},
                "具体改造项目": [],
            }
        )

        paragraph_texts = [paragraph.text.strip() for paragraph in generator.doc.paragraphs]
        self.assertIn("总体实施计划建议（三年分阶段）", paragraph_texts)
        self.assertNotIn("总体实施计划建议（两年分阶段）", paragraph_texts)

        planning_table = generator._table_by_headers(["阶段", "时间", "核心任务", "关键指标"])
        implementation_table = generator._table_by_headers(["阶段", "周期", "核心任务"])
        self.assertEqual([planning_table.cell(row, 1).text for row in range(1, 4)], ["0-12个月", "13-24个月", "25-36个月"])
        self.assertEqual([implementation_table.cell(row, 1).text for row in range(1, 4)], ["0-12个月", "13-24个月", "25-36个月"])
        self.assertNotIn("旧阶段", " ".join(cell.text for table in (planning_table, implementation_table) for row in table.rows for cell in row.cells))

    def test_legacy_two_year_fields_remain_readable_under_three_year_heading(self) -> None:
        generator = self.generator()
        generator.fill_zhigai_jianshe_fangan(
            {
                "企业名称": "测试企业",
                "建设思路": "建设思路",
                "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                "两年规划": [
                    {"阶段": "兼容阶段", "时间": "25-36个月", "核心任务": "兼容任务", "关键成果": "兼容成果"}
                ],
                "总体方案架构": [],
                "建设内容描述": [],
                "两年实施计划": [
                    {"阶段": "兼容阶段", "周期": "25-36个月", "核心任务": "兼容实施"}
                ],
                "项目预期成效": {},
                "具体改造项目": [],
            }
        )

        paragraph_texts = [paragraph.text.strip() for paragraph in generator.doc.paragraphs]
        self.assertIn("总体实施计划建议（三年分阶段）", paragraph_texts)
        planning_table = generator._table_by_headers(["阶段", "时间", "核心任务", "关键指标"])
        implementation_table = generator._table_by_headers(["阶段", "周期", "核心任务"])
        self.assertEqual(planning_table.cell(1, 0).text, "兼容阶段")
        self.assertEqual(implementation_table.cell(1, 0).text, "兼容阶段")

    def test_dynamic_projects_clone_template_numbering_hierarchy(self) -> None:
        generator = self.generator()
        generator.fill_zhigai_jianshe_fangan(
            {
                "企业名称": "测试企业",
                "建设思路": "建设思路",
                "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                "总体方案架构": [],
                "建设内容描述": [],
                "项目预期成效": {},
                "具体改造项目": [
                    {
                        "类别": "工业软件应用类",
                        "项目名称": "MES系统建设",
                        "项目内容": "建设MES",
                        "改造环节": "生产管理",
                        "建设目标": "生产透明化",
                        "预计投入": "80-100万元（估算）",
                        "投资回报周期": "18-24个月（估算）",
                        "实施步骤": "调研、实施、验收",
                    },
                    {
                        "类别": "平台与数据中台建设类",
                        "项目名称": "数据中台建设",
                        "项目内容": "建设数据中台",
                        "改造环节": "数据治理",
                        "建设目标": "数据贯通",
                        "预计投入": "80-100万元（估算）",
                        "投资回报周期": "18-24个月（估算）",
                        "实施步骤": "治理、集成、上线",
                    },
                ],
            }
        )

        paragraphs = generator.doc.paragraphs
        category = next(paragraph for paragraph in paragraphs if paragraph.text == "工业软件应用类")
        project = next(paragraph for paragraph in paragraphs if paragraph.text == "项目1：MES系统建设")
        field = next(paragraph for paragraph in paragraphs if paragraph.text == "项目内容：建设MES")
        self.assertEqual(category.style.name, "Heading 3")
        self.assertEqual(project.style.name, "Heading 4")
        self.assertEqual(field.style.name, "Normal")
        self.assertIsNotNone(category._p.pPr.numPr)
        self.assertIsNotNone(field._p.pPr.numPr)
        self.assertEqual(field._p.pPr.numPr.numId.val, 22)

    def test_full_generation_keeps_image_slots_after_all_section_edits(self) -> None:
        generator = self.generator()
        output_path = Path(self.temp_dir.name) / "full-report.docx"
        service_table_xml = generator.doc.tables[-1]._tbl.xml
        architecture_names = [
            "基础层（基础设施保障）",
            "设备与控制层（设备数采集成）",
            "平台层（平台支撑）",
            "管理应用层（企业管理应用）",
            "智能决策层（BI+AI 决策）",
        ]

        generator.generate(
            {
                "企业总体情况分段": {
                    "总体文字": "企业总体情况",
                    "主要产品列表": ["产品A"],
                    "产品工艺列表": [{"名称": "产品A", "流程": "工序A→工序B"}],
                },
                "企业总体图片": {"第一部分": [str(self.image_path)]},
                "智能工厂对标分析": {
                    "成熟度评分表已上传": True,
                    "企业名称": "测试企业",
                    "成熟度总分": "1.0",
                    "成熟度评分图片": str(self.image_path),
                    "能力子域分析": {
                        "人员域": [
                            {
                                "名称": "战略组织",
                                "得分": "0.83",
                                "现状分析": "已开展数字化规划。",
                                "差距分析": "尚未形成跨部门推进机制。",
                            },
                            {
                                "名称": "人员技能",
                                "得分": "0.54",
                                "现状分析": "具备基础操作人员。",
                                "差距分析": "缺少数据分析与系统运维能力。",
                            },
                        ]
                    },
                },
                "痛点分析": {
                    "总体概述": "总体痛点",
                    "痛点列表": [
                        {"名称": f"痛点{index}", "内容": f"痛点{index}正文"}
                        for index in range(1, 7)
                    ],
                },
                "智改数转建设方案": {
                    "企业名称": "测试企业",
                    "建设思路": "建设思路",
                    "建设目标": {"总体目标": "总体目标", "具体目标": {}},
                    "总体方案架构": [
                        {"层级": name, "内容": f"{name}内容"} for name in architecture_names
                    ],
                    "总体技术架构图": str(self.image_path),
                    "建设内容描述": [
                        {
                            "名称": name.split("（", 1)[0],
                            "建设内容": f"{name}建设内容",
                            "建设效果": f"{name}建设效果",
                            "解决痛点": f"{name}解决痛点",
                        }
                        for name in architecture_names
                    ],
                    "项目预期成效": {},
                    "具体改造项目": [],
                },
            },
            str(output_path),
        )

        document = Document(output_path)
        paragraphs = document.paragraphs
        paragraph_texts = [paragraph.text.strip() for paragraph in paragraphs]
        strategy_index = paragraph_texts.index("战略组织（得分0.83）")
        self.assertEqual(paragraph_texts[strategy_index + 1], "现状分析：已开展数字化规划。")
        self.assertEqual(paragraph_texts[strategy_index + 2], "差距分析：尚未形成跨部门推进机制。")
        self.assertIn("基础层（基础设施保障）内容", paragraph_texts)
        self.assertIn("基础层（基础设施保障）建设内容", paragraph_texts)
        self.assertIn("建设效果：基础层（基础设施保障）建设效果", paragraph_texts)
        self.assertIn("解决痛点：基础层（基础设施保障）解决痛点", paragraph_texts)
        maturity_heading = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "智能制造成熟度现状评测具体情况"
        )
        self.assertTrue(has_drawing(paragraphs[maturity_heading + 1]))
        self.assertEqual(paragraphs[maturity_heading + 2].text.strip(), "智能制造成熟度现状分析")

        final_layer = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "智能决策层（BI+AI 决策）"
        )
        architecture_caption = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "总体技术架构图"
        )
        construction_content = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "建设内容描述"
        )
        self.assertEqual(
            sum(
                has_drawing(paragraph)
                for paragraph in paragraphs[final_layer + 1 : architecture_caption]
            ),
            1,
        )
        self.assertLess(architecture_caption, construction_content)
        self.assertEqual(generator.doc.tables[-1]._tbl.xml, service_table_xml)
        self.assertEqual(
            sum(paragraph.text.strip() == "暂无。" for paragraph in paragraphs),
            2,
        )
        self.assertGreater(output_path.stat().st_size, 100_000)


if __name__ == "__main__":
    unittest.main()